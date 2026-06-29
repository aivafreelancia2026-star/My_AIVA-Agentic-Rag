"""
agents/file_editing_agent.py

File Editing Agent for Phase 2: In-App File Viewing and Editing.

Responsibilities:
  - Detect file type and decide if the file is editable.
  - Load document content in a structured form.
  - Apply safe edits (with backup and version history).
  - Produce a diff / preview before saving.
  - Save the new version.
  - Trigger re-indexing after save.
  - Report success or failure.

Security rules enforced:
  - All paths are validated against ALLOWED_ROOT so directory traversal is impossible.
  - Backups are always created before any write.
  - AI-proposed edits require an explicit confirmation flag before saving.
  - Corrupted-file writes restore the backup automatically.
"""

import os
import re
import csv
import json
import shutil
import logging
import difflib
import io
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple

logger = logging.getLogger(__name__)

# ── Allowed file types ────────────────────────────────────────────────────────

EDITABLE_TYPES   = {'.txt', '.md', '.csv', '.docx', '.xlsx'}
VIEW_ONLY_TYPES  = {'.pdf', '.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'}
ALL_SUPPORTED    = EDITABLE_TYPES | VIEW_ONLY_TYPES

# ── Safe root (all file ops are restricted to this directory) ─────────────────

def _get_allowed_root() -> Path:
    """Return the project root that all file operations must stay within."""
    try:
        from config import config as app_config
        base = Path(app_config.PROJECT_ROOT)
    except Exception:
        base = Path(__file__).parent.parent
    return base.resolve()


def _safe_path(raw: str) -> Optional[Path]:
    """
    Resolve `raw` to an absolute path and ensure it stays inside ALLOWED_ROOT.
    Returns None (and logs a warning) if the path escapes the allowed root.
    """
    allowed = _get_allowed_root()
    try:
        p = (allowed / raw.lstrip("/\\")).resolve()
    except Exception as exc:
        logger.warning("[FileEditAgent] Bad path %r: %s", raw, exc)
        return None
    if not str(p).startswith(str(allowed)):
        logger.warning("[FileEditAgent] Path traversal attempt blocked: %r → %r", raw, p)
        return None
    return p


# ── Backup / versioning ───────────────────────────────────────────────────────

def _backup_path(file_path: Path) -> Path:
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    return file_path.parent / f"{file_path.stem}_backup_{ts}{file_path.suffix}"


def _make_backup(file_path: Path) -> Optional[Path]:
    """Copy the file to a timestamped backup. Returns the backup path or None."""
    if not file_path.exists():
        return None
    bp = _backup_path(file_path)
    try:
        shutil.copy2(str(file_path), str(bp))
        logger.info("[FileEditAgent] Backup created: %s", bp)
        return bp
    except Exception as exc:
        logger.error("[FileEditAgent] Backup failed for %s: %s", file_path, exc)
        return None


def list_backups(relative_path: str) -> List[Dict[str, str]]:
    """Return all backup files for the given document, newest first."""
    fp = _safe_path(relative_path)
    if fp is None:
        return []
    pattern = f"{fp.stem}_backup_*{fp.suffix}"
    backups = sorted(fp.parent.glob(pattern), key=lambda p: p.stat().st_mtime, reverse=True)
    return [{"name": b.name, "path": str(b.relative_to(_get_allowed_root()))} for b in backups]


def restore_backup(backup_relative: str, target_relative: str) -> Dict[str, Any]:
    """Overwrite target with the chosen backup."""
    src = _safe_path(backup_relative)
    tgt = _safe_path(target_relative)
    if src is None or tgt is None:
        return {"success": False, "error": "Invalid path."}
    if not src.exists():
        return {"success": False, "error": "Backup file not found."}
    try:
        shutil.copy2(str(src), str(tgt))
        return {"success": True, "message": f"Restored from {src.name}"}
    except Exception as exc:
        return {"success": False, "error": str(exc)}


# ── File info ─────────────────────────────────────────────────────────────────

def get_file_info(relative_path: str) -> Dict[str, Any]:
    """Return metadata for a single file."""
    fp = _safe_path(relative_path)
    if fp is None:
        return {"error": "Invalid path."}
    if not fp.exists():
        return {"error": "File not found."}

    suffix = fp.suffix.lower()
    stat = fp.stat()
    return {
        "name":         fp.name,
        "path":         relative_path,
        "suffix":       suffix,
        "size_bytes":   stat.st_size,
        "modified":     datetime.fromtimestamp(stat.st_mtime).isoformat(),
        "is_editable":  suffix in EDITABLE_TYPES,
        "is_viewable":  suffix in ALL_SUPPORTED,
    }


# ── Read / view ───────────────────────────────────────────────────────────────

def read_file_content(relative_path: str) -> Dict[str, Any]:
    """
    Load a file and return its content in a structured way.

    Returns a dict with at least: {"type": ..., "content": ..., "editable": bool}
    """
    fp = _safe_path(relative_path)
    if fp is None:
        return {"error": "Invalid path."}
    if not fp.exists():
        return {"error": "File not found."}

    suffix = fp.suffix.lower()

    # ── Plain text / markdown ──────────────────────────────────────────────────
    if suffix in {'.txt', '.md'}:
        try:
            text = fp.read_text(encoding='utf-8', errors='replace')
            return {"type": "text", "content": text, "editable": True, "encoding": "utf-8"}
        except Exception as exc:
            return {"error": f"Cannot read file: {exc}"}

    # ── CSV ───────────────────────────────────────────────────────────────────
    if suffix == '.csv':
        try:
            rows = []
            with open(fp, newline='', encoding='utf-8', errors='replace') as f:
                reader = csv.reader(f)
                for row in reader:
                    rows.append(row)
            return {"type": "csv", "rows": rows, "editable": True}
        except Exception as exc:
            return {"error": f"Cannot read CSV: {exc}"}

    # ── DOCX ─────────────────────────────────────────────────────────────────
    if suffix == '.docx':
        try:
            import docx
            doc = docx.Document(str(fp))
            paragraphs = [{"text": p.text, "style": p.style.name} for p in doc.paragraphs]
            tables = []
            for tbl in doc.tables:
                table_data = [[cell.text for cell in row.cells] for row in tbl.rows]
                tables.append(table_data)
            return {"type": "docx", "paragraphs": paragraphs, "tables": tables, "editable": True}
        except ImportError:
            return {"error": "python-docx not installed. Run: pip install python-docx"}
        except Exception as exc:
            return {"error": f"Cannot read DOCX: {exc}"}

    # ── XLSX ─────────────────────────────────────────────────────────────────
    if suffix == '.xlsx':
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(fp), read_only=True, data_only=True)
            sheets = {}
            for name in wb.sheetnames:
                ws = wb[name]
                rows = []
                for row in ws.iter_rows(values_only=True):
                    rows.append([str(c) if c is not None else "" for c in row])
                sheets[name] = rows
            return {"type": "xlsx", "sheets": sheets, "sheet_names": wb.sheetnames, "editable": True}
        except ImportError:
            return {"error": "openpyxl not installed. Run: pip install openpyxl"}
        except Exception as exc:
            return {"error": f"Cannot read XLSX: {exc}"}

    # ── PDF ───────────────────────────────────────────────────────────────────
    if suffix == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(str(fp))
            pages = []
            for i, page in enumerate(doc):
                pages.append({"page": i + 1, "text": page.get_text("text")})
            doc.close()
            return {"type": "pdf", "pages": pages, "editable": False}
        except ImportError:
            return {"type": "pdf", "pages": [], "editable": False,
                    "note": "PyMuPDF not installed; cannot extract text."}
        except Exception as exc:
            return {"error": f"Cannot read PDF: {exc}"}

    # ── Image ─────────────────────────────────────────────────────────────────
    if suffix in {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'}:
        return {"type": "image", "path": relative_path, "editable": False}

    return {"error": f"Unsupported file type: {suffix}"}


# ── Write / save ──────────────────────────────────────────────────────────────

def save_text_file(
    relative_path: str,
    new_content: str,
    confirmed: bool = False,
    save_as_new_version: bool = False,
) -> Dict[str, Any]:
    """
    Save updated text (.txt / .md) content.
    Always creates a backup. Requires confirmed=True to actually write.
    """
    fp = _safe_path(relative_path)
    if fp is None:
        return {"success": False, "error": "Invalid path."}
    if fp.suffix.lower() not in {'.txt', '.md'}:
        return {"success": False, "error": "Use save_text_file only for .txt / .md files."}

    old_content = fp.read_text(encoding='utf-8', errors='replace') if fp.exists() else ""
    diff = list(difflib.unified_diff(
        old_content.splitlines(keepends=True),
        new_content.splitlines(keepends=True),
        fromfile="original",
        tofile="modified",
        n=3,
    ))

    if not confirmed:
        return {
            "success": False,
            "pending_confirmation": True,
            "diff": "".join(diff[:100]),
            "message": "Preview generated. Send confirmed=True to save.",
        }

    backup = _make_backup(fp) if fp.exists() else None

    target = fp if not save_as_new_version else _version_path(fp)
    try:
        target.write_text(new_content, encoding='utf-8')
        return {
            "success": True,
            "saved_to": str(target.relative_to(_get_allowed_root())),
            "backup": str(backup.relative_to(_get_allowed_root())) if backup else None,
            "diff": "".join(diff[:100]),
        }
    except Exception as exc:
        if backup and backup.exists():
            shutil.copy2(str(backup), str(fp))
        return {"success": False, "error": str(exc)}


def save_csv_file(
    relative_path: str,
    rows: List[List[str]],
    confirmed: bool = False,
    save_as_new_version: bool = False,
) -> Dict[str, Any]:
    """Save updated CSV content (rows is a list of lists)."""
    fp = _safe_path(relative_path)
    if fp is None:
        return {"success": False, "error": "Invalid path."}
    if fp.suffix.lower() != '.csv':
        return {"success": False, "error": "File is not a CSV."}

    if not confirmed:
        preview = rows[:5]
        return {
            "success": False,
            "pending_confirmation": True,
            "preview_rows": preview,
            "total_rows": len(rows),
            "message": "Preview generated. Send confirmed=True to save.",
        }

    backup = _make_backup(fp) if fp.exists() else None
    target = fp if not save_as_new_version else _version_path(fp)
    try:
        with open(target, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerows(rows)
        return {
            "success": True,
            "saved_to": str(target.relative_to(_get_allowed_root())),
            "backup": str(backup.relative_to(_get_allowed_root())) if backup else None,
        }
    except Exception as exc:
        if backup and backup.exists():
            shutil.copy2(str(backup), str(fp))
        return {"success": False, "error": str(exc)}


def save_docx_file(
    relative_path: str,
    paragraphs: List[Dict[str, str]],
    confirmed: bool = False,
    save_as_new_version: bool = True,
) -> Dict[str, Any]:
    """
    Save updated DOCX paragraphs.
    Defaults to save_as_new_version=True to protect the original.
    """
    fp = _safe_path(relative_path)
    if fp is None:
        return {"success": False, "error": "Invalid path."}
    if fp.suffix.lower() != '.docx':
        return {"success": False, "error": "File is not a DOCX."}

    if not confirmed:
        return {
            "success": False,
            "pending_confirmation": True,
            "paragraph_count": len(paragraphs),
            "message": "Send confirmed=True to save. Defaults to a new version to protect original.",
        }

    try:
        import docx as docx_lib
    except ImportError:
        return {"success": False, "error": "python-docx not installed."}

    backup = _make_backup(fp) if fp.exists() else None
    target = fp if not save_as_new_version else _version_path(fp)

    try:
        doc = docx_lib.Document(str(fp)) if fp.exists() else docx_lib.Document()
        # Overwrite paragraphs in-place (best-effort style preservation)
        for i, para_data in enumerate(paragraphs):
            text = para_data.get("text", "")
            if i < len(doc.paragraphs):
                doc.paragraphs[i].text = text
            else:
                doc.add_paragraph(text)
        doc.save(str(target))
        return {
            "success": True,
            "saved_to": str(target.relative_to(_get_allowed_root())),
            "backup": str(backup.relative_to(_get_allowed_root())) if backup else None,
            "warning": "Complex formatting (tables, images) may not be preserved in the edited paragraphs.",
        }
    except Exception as exc:
        if backup and backup.exists():
            shutil.copy2(str(backup), str(fp))
        return {"success": False, "error": str(exc)}


def save_xlsx_file(
    relative_path: str,
    sheet_name: str,
    rows: List[List[str]],
    confirmed: bool = False,
    save_as_new_version: bool = True,
) -> Dict[str, Any]:
    """
    Save updated cell data for one sheet in an XLSX workbook.
    Other sheets are preserved unchanged.
    """
    fp = _safe_path(relative_path)
    if fp is None:
        return {"success": False, "error": "Invalid path."}
    if fp.suffix.lower() != '.xlsx':
        return {"success": False, "error": "File is not an XLSX."}

    if not confirmed:
        return {
            "success": False,
            "pending_confirmation": True,
            "sheet": sheet_name,
            "preview_rows": rows[:5],
            "total_rows": len(rows),
            "message": "Send confirmed=True to save. Formulas and styles may be lost.",
            "warning": "Complex formulas and cell styles may not be preserved.",
        }

    try:
        import openpyxl
    except ImportError:
        return {"success": False, "error": "openpyxl not installed."}

    backup = _make_backup(fp) if fp.exists() else None
    target = fp if not save_as_new_version else _version_path(fp)

    try:
        wb = openpyxl.load_workbook(str(fp))
        if sheet_name not in wb.sheetnames:
            return {"success": False, "error": f"Sheet '{sheet_name}' not found."}
        ws = wb[sheet_name]
        # Clear the sheet and rewrite rows
        for row in ws.iter_rows():
            for cell in row:
                cell.value = None
        for r_idx, row_data in enumerate(rows, start=1):
            for c_idx, val in enumerate(row_data, start=1):
                ws.cell(row=r_idx, column=c_idx, value=val)
        wb.save(str(target))
        return {
            "success": True,
            "saved_to": str(target.relative_to(_get_allowed_root())),
            "backup": str(backup.relative_to(_get_allowed_root())) if backup else None,
            "warning": "Formulas and cell formatting may have been simplified.",
        }
    except Exception as exc:
        if backup and backup.exists():
            shutil.copy2(str(backup), str(fp))
        return {"success": False, "error": str(exc)}


# ── Re-indexing ───────────────────────────────────────────────────────────────

def reindex_file(relative_path: str) -> Dict[str, Any]:
    """
    Remove old FAISS chunks for the file and re-process it into the vector store.

    This function locates the unified document processor and triggers it for
    the specific file, then reloads the doc_manager.
    """
    fp = _safe_path(relative_path)
    if fp is None:
        return {"success": False, "error": "Invalid path."}
    if not fp.exists():
        return {"success": False, "error": "File not found."}

    try:
        from unified_document_processor import MultiFormatDocumentProcessor
        from config import config as _cfg

        processor = MultiFormatDocumentProcessor()
        processor.process_file(str(fp), _cfg.EMBEDDINGS_DIR)

        # Reload doc_manager so new chunks are live
        try:
            from app_state import get_system_state
            _, _, dm, _, _, _, _ = get_system_state()
            if dm:
                dm.load_all_documents()
        except Exception as reload_exc:
            logger.warning("[FileEditAgent] doc_manager reload warning: %s", reload_exc)

        return {"success": True, "message": f"Re-indexed: {fp.name}"}
    except ImportError as ie:
        return {
            "success": False,
            "error": f"unified_document_processor not available: {ie}",
        }
    except Exception as exc:
        return {"success": False, "error": str(exc)}


# ── Edit log ──────────────────────────────────────────────────────────────────

_EDIT_LOG_FILE = Path(__file__).parent.parent / "edit_log.json"


def append_edit_log(entry: Dict[str, Any]) -> None:
    """Append an entry to the edit log (JSON Lines)."""
    try:
        with open(_EDIT_LOG_FILE, 'a', encoding='utf-8') as f:
            f.write(json.dumps(entry, default=str) + "\n")
    except Exception as exc:
        logger.warning("[FileEditAgent] Could not write edit log: %s", exc)


def get_edit_log(last_n: int = 50) -> List[Dict[str, Any]]:
    """Return the last N edit log entries."""
    if not _EDIT_LOG_FILE.exists():
        return []
    entries = []
    try:
        with open(_EDIT_LOG_FILE, encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line:
                    try:
                        entries.append(json.loads(line))
                    except Exception:
                        pass
    except Exception:
        pass
    return entries[-last_n:]


# ── Helpers ───────────────────────────────────────────────────────────────────

def _version_path(fp: Path) -> Path:
    """Generate a timestamped 'new version' path."""
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    return fp.parent / f"{fp.stem}_v{ts}{fp.suffix}"


def apply_ai_edits(
    relative_path: str,
    instruction: str,
    llm=None,
    confirmed: bool = False,
) -> Dict[str, Any]:
    """
    Apply an AI-generated edit to a text or markdown file.

    Steps:
      1. Read current content.
      2. Ask the LLM to produce the new version (if llm is provided).
      3. Return a diff for preview.
      4. If confirmed=True, save via save_text_file().

    If llm is None, returns an error asking the user to supply the new content.
    """
    suffix = Path(relative_path).suffix.lower()
    if suffix not in {'.txt', '.md'}:
        return {"success": False, "error": "AI edits are currently supported only for .txt and .md files."}

    content_result = read_file_content(relative_path)
    if "error" in content_result:
        return {"success": False, "error": content_result["error"]}

    old_content = content_result.get("content", "")

    if llm is None:
        return {
            "success": False,
            "error": "No LLM available for AI editing. Supply the new content directly via save_text_file().",
        }

    prompt = (
        f"You are a document editor. Apply the following instruction to the document below.\n"
        f"Return ONLY the full updated document text — no explanation, no markdown fences.\n\n"
        f"INSTRUCTION: {instruction}\n\n"
        f"DOCUMENT:\n{old_content}\n\nUPDATED DOCUMENT:"
    )
    try:
        response = llm.invoke(prompt)
        new_content = response.content if hasattr(response, 'content') else str(response)
    except Exception as exc:
        return {"success": False, "error": f"LLM error: {exc}"}

    if not confirmed:
        diff = list(difflib.unified_diff(
            old_content.splitlines(keepends=True),
            new_content.splitlines(keepends=True),
            fromfile="original",
            tofile="ai_edited",
            n=3,
        ))
        return {
            "success": False,
            "pending_confirmation": True,
            "proposed_content": new_content,
            "diff": "".join(diff[:120]),
            "message": "AI edit preview ready. Send confirmed=True with proposed_content to save.",
        }

    return save_text_file(relative_path, new_content, confirmed=True)
